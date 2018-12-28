import docx
import os
import subprocess
import pandas as pd
import pickle
from pathlib import PurePath

def crawler_arquivos_tipo(extension, dir_ = os.getcwd()):
    '''Percorre uma árvore de diretórios e obtém os 
    paths absolutos de todos os arquivos de uma determinada extensão'''
    #Essa função está fora da classe porque ela pertence a uma outra classe que abstrai diversas funções de parseamento de arquivos.
    #Coloquei ela aqui por comodidade, para não ter que realizar o import
    
    paths = [(tupla[0], tupla[2]) for tupla in list(os.walk(dir_))]
    path_files = [os.path.abspath(os.path.join(tupla[0], file)) for tupla in paths for file in tupla[1]]
    files = [path for path in path_files if os.path.isfile(path) and extension == PurePath(path).suffix]
    
    return files

class ParserDocx():
    '''Implementa o pipeline de conversão de parseamento dos arquivos .doc que contém
    o cadastro de expedientes sobre os Croquis Patrimominais de CGPATRI.
    *** Caso queira instanciar o objeto utilizando a função main, o DataFrame final
        pode ser acessado por meio do atributo .df do objeto'''
    
    def __init__(self, dir_arquivos_doc, dir_salvar_docx):
        
        self.dir_docs = dir_arquivos_doc
        self.dir_docx = dir_salvar_docx
    
    def __main__(self):
        
        self.n_cvt = self.pipe_conversao_docx(self.dir_docx)
        print('Conversao finalizada. Iniciando o parseamento')
        self.df = self.parser_docs()
    
    def doc_to_docx(self, doc_path):
        '''Gera um subprocesso que chama o LibreOffice sem abrir a GUI e executa a funcionalidade
        de conversão do arquivo .doc para um arquivo .docx'''
    
        subprocess.call([r'C:\Program Files\LibreOffice\program\soffice.exe', '--convert-to', 'docx', '"{}"'.format(doc_path)],
                     shell = True)
        
    def doc_to_docx_todos(self, paths):
        '''Recebe uma lista de paths de arquivos .doc e converte todos
        eles para .docx'''
        
        n_comput = 0
        for f_path in paths:

            self.doc_to_docx(f_path)
            n_comput +=1
        return n_comput
    
    def nao_convertidos(self, docs, docs_x):
        '''Recebe a lista de paths de arquivos .doc original e compara
        com a lista de paths dos arquivos .docx convertidos, devolvendo
        a lista de arquivos não-convertidos'''
    
        docs_x = [os.path.basename(item) for item in docs_x]
        docs = [''.join([os.path.basename(item), 'x']) for item in docs]

        return [item for item in docs if item not in docs_x] 
    
    def pipe_conversao_docx(self, dir_docs = self.dir_docs, dir_docx = self.dir_docx):
        '''Implementa o pipeline de conversão dos arquivos .doc para doc.x,
        parseando todos os arquivos de um dado diretorio'''
    
        docs = crawler_arquivos_tipo('.doc', dir_ = dir_docs)
        print('Total de documentos a converter: ', len(docs))
        cwd = os.getcwd()
        os.mkdir(self.dir_docx)
        os.chdir(self.dir_docx)
        n_comput = self.doc_to_docx_todos(docs)
        print('Foram realizadas {} computacoes'.format(n_comput))
        docx = crawler_arquivos_tipo('.docx')
        n_cvt = self.nao_convertidos(docs, docx)
        os.chdir(cwd)

        return n_cvt
    
    def pickles_nao_convertidos(self, nome_pickles = 'nao_convertidos.pi', dir_alvo = ''):
        '''Chama o pipeline de conversão de arquivos e salva os paths dos arquivos não convertidos em
        um pickles'''
        
        n_cvt = self.pipe_conversao_docx(self.dir_docx)
        if dir_alvo:
            cwd = os.getcwd()
            os.chdir(dir_alvo)
        with open(nome_pickles, 'wb') as f:
            pickle.dumps(n_cvt)
        os.chdir(cwd)
    
    def mapear_dados_tabela(self, docx_, table_num):
        '''Abstrai o mapeamento de dados de uma tabela em um arquivo 
        .docx'''
        
        dados = {}
        for i, row in enumerate(docx_.tables[table_num].rows):
            for z, cell in enumerate(row.cells):
                if i not in dados.keys():
                    dados[i] =  {}
                try:
                    txt = cell.text.strip()
                    dados[i][z] = txt
                except Exception as e:
                    print(e)
                    dados[i][z] = ''
        return dados
    
    def num_plan(self, file):
        '''Puxa o número da planilha do arquivo .docx'''
        
        try:
            return file.paragraphs[0].text.strip()
        except:
            return ''
        
    def dt_tramit(self, file):
        '''Puxa o campo da data de tramitação do arquivo .docx'''
        
        try:
            return file.paragraphs[5].text.strip().replace('Data de Tramitação:', '')
        except: 
            return '
    def table_0(self, file, dados):
        '''Parseia a primeira tabela do arquivo .docx'''
    
        dados_tb = self.mapear_dados_tabela(file, 0)
        dados['Croquis'] = dados_tb[0][1]
        dados['Área'] = dados_tb[0][3]
        dados['Processo'] = dados_tb[1][1]
        dados['TID'] = dados_tb[2][1]
        dados['Expediente'] = dados_tb[3][1]
        dados['Interessado'] = dados_tb[4][1]
        dados['Assunto'] = dados_tb[5][1]
        dados['Local'] = dados_tb[6][1]    
    
    def table_1(self, file, dados):
        '''Parseia a segunda tabela do arquivo .docx'''
    
        dados_tb = mapear_dados_tabela(file, 1)
        if 'x' in dados_tb[0][1].lower():
            dados['Anotação'] = 1
        else:
            dados['Anotação'] = 0
        if 'x' in dados_tb[0][3].lower():
            dados['Informação'] = 1
        else:
            dados['Informação'] = 0
            
    def table_2(self, file, dados):
        '''Parseia a terceira tabela do arquivo .docx'''
    
        dados_tb = self.mapear_dados_tabela(file, 2)
        dados['Despacho'] = dados_tb[0][1]
        dados['DOM'] = dados_tb[1][1]
        
    def table_3(self, file, dados):
        '''Parseia a quarta tabela do arquivo .docx'''
    
        dados_tb = self.mapear_dados_tabela(file, 3)
        dados['Observação/Vistorias'] = dados_tb[0][0].replace('Observação/Vistorias:', '')
        
    def table_4(self, file, dados):
        '''Parseia a quinta (última) tabela do arquivo .docx'''
    
        dados_tb = self.mapear_dados_tabela(file, 4)
        dados['Data'] = dados_tb[0][1]
        dados['Autor'] = dados_tb[0][2].replace('Nome:', '')
        
    def mapear_dados_file(self, file):
        '''Parseia todos os dados do arquivo .docx'''
    
        dados = {}
        dados['Planilha_num'] = self.num_plan(file)
        dados['Data_Tramit'] = self.dt_tramit(file)
        self.table_0(file, dados)
        self.table_1(file, dados)
        self.table_2(file, dados)
        self.table_3(file, dados)
        self.table_4(file, dados)

        return dados
    
    def parser_docs(self, dir_docx = self.dir_docx):
        '''Parseia todos os arquivos .docx e retorna um DataFrame com os dados'''
    
        dados = []
        files = crawler_arquivos_tipo('.docx', dir_ = dir_docx)
        for file in files:
            try:
                file = docx.Document(file)
                dados.append(self.mapear_dados_file(file))
            except:
                pass
        print((len(files) - len(dados))/len(files))
        return pd.DataFrame(dados)
        
